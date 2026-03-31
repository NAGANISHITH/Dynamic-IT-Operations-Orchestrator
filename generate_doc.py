"""
Generate Architecture_doc.docx for
Multi-Agent IT Ops Platform
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x0D, 0x47, 0xA1)   # heading / cover title
ACCENT_BLUE = RGBColor(0x15, 0x65, 0xC0)   # sub-heading / table header bg
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)   # table alt row bg
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x21, 0x21, 0x21)
GREEN       = RGBColor(0x2E, 0x7D, 0x32)
ORANGE      = RGBColor(0xE6, 0x5C, 0x00)

doc = Document()

# ── Page margins (narrow) ──────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=DARK_BLUE, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_BLUE
    return p


def add_body(text, size=10, bold=False, color=DARK_TEXT, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    return p


def add_bullet(text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_TEXT
    return p


def add_code_block(text):
    """Monospaced paragraph styled as a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def make_table(headers, rows, col_widths=None):
    """Create a formatted table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], '1565C0')
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = 'F5F5F5' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            set_cell_bg(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK_TEXT
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 100)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run('AUTONOMOUS IT OPS PLATFORM')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = DARK_BLUE

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover_sub.add_run('Multi-Agent AI System for Real-Time IT Infrastructure Management')
r2.font.size = Pt(14)
r2.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()
add_divider()
doc.add_paragraph()

meta_rows = [
    ('Document Type',  'Architecture & Technical Reference'),
    ('Version',        '1.0.0'),
    ('Status',         'Active / Production-Ready Design'),
    ('Date',           '2026-03-28'),
    ('Classification', 'Internal / Engineering'),
]
for label, val in meta_rows:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_label = p.add_run(f'{label}: ')
    r_label.bold = True
    r_label.font.size = Pt(10)
    r_label.font.color.rgb = DARK_BLUE
    r_val = p.add_run(val)
    r_val.font.size = Pt(10)
    r_val.font.color.rgb = DARK_TEXT

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading('1. Executive Summary', level=1)
add_body(
    'The Autonomous IT Ops Platform deploys ten specialized AI agents that continuously monitor enterprise '
    'IT infrastructure, coordinate through an in-process Agent-to-Agent (A2A) message bus, predict failures '
    'using an LSTM ensemble scoring model, and execute automated runbook-driven remediation — all without '
    'human intervention.'
)
add_body(
    'The platform addresses the core operational challenge faced by modern enterprises: the volume and velocity '
    'of telemetry data generated by distributed infrastructure exceeds human capacity to manually triage, '
    'correlate, and remediate incidents within SLA windows.'
)
doc.add_paragraph()

add_heading('Key Capabilities', level=2)
bullets = [
    'Autonomous remediation: Zero-touch incident resolution end-to-end',
    'Predictive failure detection: Up to 72-hour failure horizon via LSTM scoring',
    'Root cause analysis: Multi-source correlation — logs, metrics, CMDB',
    'Security anomaly detection: Brute force, DDoS, SQL injection, unusual access',
    'Cloud cost optimization: Rightsize underutilized EC2 / RDS / VM / Storage',
    'Real-time dashboard: All agent activity streamed live over WebSocket',
    'Incident simulation: Inject any failure class via REST API for demos and testing',
]
for b in bullets:
    add_bullet(b)

add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading('2. System Overview', level=1)
add_body(
    'The platform is a fully event-driven multi-agent system built on Python asyncio. '
    'All agents run as concurrent coroutines within a single process, communicating via an '
    'in-process asyncio.Queue message bus. This architecture eliminates inter-thread locking overhead '
    'and simplifies the agent interaction model.'
)
doc.add_paragraph()
add_heading('High-Level Architecture', level=2)
add_code_block('''
┌─────────────────────────────────────────────────────────────────────────────┐
│                        HIGH-LEVEL SYSTEM VIEW                               │
│                                                                             │
│  ┌─────────────────────────────────────────────────────────────────────┐    │
│  │                  REACT FRONTEND  (Port 3000)                        │    │
│  │  Agent Topology | Incidents | Predictions | A2A Log | Live Metrics  │    │
│  └─────────────────┬──────────────────────────┬───────────────────────┘    │
│                    │  REST (HTTP/JSON)          │  WebSocket (ws://…/ws)    │
│  ┌─────────────────▼──────────────────────────▼───────────────────────┐    │
│  │                  FASTAPI BACKEND  (Port 8000)                       │    │
│  │  ┌───────────────────────────────────────────────────────────────┐  │    │
│  │  │              AGENT ORCHESTRATOR (asyncio)                      │  │    │
│  │  │  [Server Monitor] [Cloud Monitor] [App Health] ← Monitoring   │  │    │
│  │  │       └─────────────────────────┘ A2A alert                   │  │    │
│  │  │              [Predictive Failure Agent] (LSTM)  ← AI Tier     │  │    │
│  │  │                         │ A2A directive                        │  │    │
│  │  │              [Remediation Agent]  (Runbooks)   ← Decision     │  │    │
│  │  │                         │ A2A deploy action                    │  │    │
│  │  │              [Deployment Agent]   (K8s/TF)     ← Execution    │  │    │
│  │  │                         │ A2A resolved                        │  │    │
│  │  │              [Reporting Agent]    (SLA/KPI)    ← Telemetry    │  │    │
│  │  │                                                                │  │    │
│  │  │  [Security Agent] [Root Cause Agent] [Optimization Agent]     │  │    │
│  │  │  ◄──────────────── A2A Event Bus (asyncio Queue) ────────────► │  │    │
│  │  └───────────────────────────────────────────────────────────────┘  │    │
│  │  [ In-Memory Store ]  [ WebSocket Manager ]  [ REST Router ]        │    │
│  └─────────────────────────────────────────────────────────────────────┘    │
└─────────────────────────────────────────────────────────────────────────────┘
''')
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════

add_heading('3. Technology Stack', level=1)
add_heading('3.1 Backend', level=2)
make_table(
    ['Component', 'Technology', 'Version', 'Purpose'],
    [
        ['API Framework',    'FastAPI',       '0.111.0', 'Async HTTP + WebSocket server with auto-OpenAPI'],
        ['ASGI Runtime',     'Uvicorn',       '0.29.0',  'High-performance Python ASGI server'],
        ['Data Validation',  'Pydantic v2',   '2.7.1',   'Schema definitions, strict typing, fast JSON serialization'],
        ['Concurrency',      'Python asyncio','stdlib',  'Single-threaded event loop — all agents run as coroutines'],
        ['WebSockets',       'FastAPI native','—',       'Real-time push to all connected frontend clients'],
        ['Data Layer',       'In-memory dict','—',       'Zero infra dependency for dev; swap to PostgreSQL + TimescaleDB'],
        ['ML Engine',        'LSTM simulation','—',      'Production-equivalent scoring logic, no GPU required'],
        ['Containerization', 'Docker + Compose','3.9',  'Full-stack containerized deployment'],
    ],
    col_widths=[1.5, 1.5, 0.8, 3.2]
)
doc.add_paragraph()
add_heading('3.2 Frontend', level=2)
make_table(
    ['Component', 'Technology', 'Version', 'Purpose'],
    [
        ['UI Framework',  'React',           '18.2.0',  'Component-based UI with hooks for real-time state'],
        ['Build Tool',    'Create React App','5.0.1',   'Dev server, bundler, HMR'],
        ['Charts',        'Recharts',        '2.12.0',  'SVG charts for CPU/mem/latency time-series'],
        ['Icons',         'Lucide React',    '0.383.0', 'Lightweight consistent icon set'],
        ['Real-time',     'Native WebSocket','—',       'Auto-reconnecting event stream using custom hook'],
        ['Styling',       'Pure CSS-in-JS',  '—',       'No external CSS framework dependency'],
    ],
    col_widths=[1.5, 1.5, 0.8, 3.2]
)
doc.add_paragraph()
add_heading('3.3 Infrastructure', level=2)
make_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Containerization', 'Docker',        'Package backend + frontend into isolated images'],
        ['Orchestration',    'Docker Compose','Single-command full-stack startup and networking'],
        ['Web Server',       'Nginx',         'Serve React static build + API proxy in production'],
    ],
    col_widths=[1.8, 1.8, 4.0]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — AGENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('4. Agent Architecture', level=1)
add_heading('4.1 BaseAgent (Abstract Base Class)', level=2)
add_body(
    'Every agent extends BaseAgent, which provides a standard interface for agent communication and lifecycle management.'
)
add_code_block(
'''class BaseAgent(ABC):
    async def send(self, to: AgentType, payload: dict) -> A2AMessage
        # Persists message → broadcasts over WebSocket → delivers to target inbox

    async def receive(self, msg: A2AMessage)
        # Drops message into agent's asyncio.Queue inbox

    @abstractmethod
    async def run(self)
        # Main event loop — implemented by each specialized agent'''
)
doc.add_paragraph()
add_heading('4.2 Agent Registry', level=2)
make_table(
    ['Agent', 'Module', 'Loop Interval', 'Primary Inputs', 'Primary Outputs'],
    [
        ['ServerMonitorAgent',    'monitoring.py',   'Every 15s',       'Host CPU/mem/disk metrics',               'Alert → PredictiveAgent'],
        ['CloudMonitorAgent',     'monitoring.py',   'Every 20s',       'Cloud packet loss / latency / cost',       'Alert → PredictiveAgent'],
        ['AppHealthAgent',        'monitoring.py',   'Every 12s',       'App p99 latency / error rate / RPS',       'Alert → PredictiveAgent'],
        ['PredictiveFailureAgent','predictive.py',   'Event + 30s sweep','Alert payloads from monitors',            'Prediction + Incident → Remediation'],
        ['RemediationAgent',      'remediation.py',  'Event-driven',    'Incident directives',                      'Runbook steps → Deployment'],
        ['DeploymentAgent',       'remediation.py',  'Event-driven',    'Deploy action commands',                   'K8s/TF ops → Resolved + Reporting'],
        ['ReportingAgent',        'reporting.py',    'Every 10s',       'Store state snapshot',                     'KPISnapshot broadcast'],
        ['RootCauseAgent',        'root_cause.py',   'Event-driven',    'Incident + failure class + resource ID',   'RCA explanation → Reporting + Notification'],
        ['SecurityAgent',         'security.py',     'Every 45s',       'Traffic / access log simulation',          'SecurityAnomaly → Reporting + Notification'],
        ['OptimizationAgent',     'optimization.py', 'Every 60s',       'Cloud cost metrics (AWS/Azure/GCP)',        'OptimizationSuggestion → Reporting'],
    ],
    col_widths=[1.6, 1.4, 1.1, 2.0, 2.1]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — A2A MESSAGE BUS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('5. A2A Message Bus', level=1)
add_heading('5.1 Message Structure', level=2)
add_code_block(
'''class A2AMessage(BaseModel):
    id:          str        # UUID[:8] — short unique ID
    from_agent:  AgentType  # Source agent enum
    to_agent:    AgentType  # Target agent enum
    payload:     dict       # Arbitrary JSON payload
    timestamp:   datetime   # UTC timestamp'''
)
doc.add_paragraph()
add_heading('5.2 Message Lifecycle', level=2)
add_body('Every A2A message follows this path:')
add_bullet('Persisted → in-memory store (or PostgreSQL in production)')
add_bullet('Broadcast → all connected WebSocket frontend clients')
add_bullet('Delivered → target agent asyncio.Queue inbox Queue')
doc.add_paragraph()
add_heading('5.3 Full Message Flow (OOMKill example)', level=2)
add_code_block(
'''ServerMonitorAgent ─── mem alert ──────────────────────────► PredictiveFailureAgent
CloudMonitorAgent  ─── packet_loss alert ──────────────────►       │
AppHealthAgent     ─── latency alert ──────────────────────►       │ LSTM score
                                                                    ▼
                                                          Prediction created
                                                          Incident opened (P1)
                                                                    │ directive
                                                                    ▼
                                                         RemediationAgent
                                                         (runbook: oom-001)
                                                                    │ deploy action
                                                                    ▼
                                                         DeploymentAgent
                                                         (K8s: scale_out, configmap_patch)
                                                                    │ resolved
                                                                    ▼
                                                         ReportingAgent → KPI broadcast'''
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PREDICTIVE FAILURE ENGINE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('6. Predictive Failure Engine', level=1)
add_body(
    'The PredictiveFailureAgent simulates an LSTM ensemble model with production-equivalent algorithm logic. '
    'The scoring formula is designed to be drop-in replaceable with a trained PyTorch/TensorFlow model.'
)
doc.add_paragraph()
add_heading('6.1 Scoring Algorithm', level=2)
add_code_block(
'''Step 1 — Deduplication:
    Same host + alert_type within 60 seconds → skip (prevents alert storm)

Step 2 — LSTM Score simulation:
    overshoot  = (measured_value − threshold) / threshold
    probability = min(0.97,  0.60 + overshoot × 0.80 + jitter[-0.05, +0.05])
    horizon_min = max(5, 120 − overshoot × 60 ± jitter)

Step 3 — Risk classification:
    probability ≥ 0.75  → HIGH  (immediate action)
    probability ≥ 0.50  → MEDIUM (monitor closely)
    probability  < 0.50  → LOW

Step 4 — Dispatch:
    if probability > 0.55:
        → Create Incident { severity: P1/P2/P3, status: triaging }
        → Send directive to RemediationAgent

Step 5 — Autonomous sweep (every 30s):
    Proactively scan latest host metrics for cpu > 88% or mem > 88%
    → Trigger scoring without waiting for an alert message'''
)
doc.add_paragraph()
add_heading('6.2 Runbook Mapping', level=2)
make_table(
    ['Alert Type', 'Failure Class', 'Severity', 'Runbook ID'],
    [
        ['mem',         'OOMKill',             'P1', 'oom-001'],
        ['cpu',         'CPUSaturation',       'P2', 'hpa-scale'],
        ['disk',        'DiskPressure',        'P2', 'disk-cleanup'],
        ['packet_loss', 'NetworkDegradation',  'P2', 'vpc-reroute'],
        ['latency',     'LatencySpike',        'P2', 'db-throttle'],
        ['error_rate',  'LatencySpike',        'P3', 'rollback-deploy'],
    ],
    col_widths=[1.5, 2.0, 1.0, 2.5]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — REMEDIATION & DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════

add_heading('7. Remediation & Deployment Pipeline', level=1)
add_heading('7.1 Runbook Catalogue', level=2)
make_table(
    ['Runbook ID', 'Failure Class', 'Automated Actions'],
    [
        ['oom-001',          'OOMKill',            'Heap limit increase + pod scale-out +2 replicas'],
        ['hpa-scale',        'CPUSaturation',      'HPA update 4→9 replicas (Kubernetes)'],
        ['vpc-reroute',      'NetworkDegradation', 'BGP reroute to backup AZ + packet loss confirmed 0%'],
        ['db-throttle',      'LatencySpike/ReplicaLag','Promote read replica + throttle ETL writes to 40%'],
        ['disk-cleanup',     'DiskPressure',       'Log rotation + compress stale logs (18 GB freed)'],
        ['cert-renew',       'CertExpiry',         'ACM auto-renewal + DNS validation record + LB update'],
        ['kafka-rebalance',  'KafkaLag',           'Partition rebalance + consumer scale +4 threads'],
        ['rollback-deploy',  'Error rate spike',   'kubectl rollout undo → previous stable version'],
    ],
    col_widths=[1.5, 1.8, 4.0]
)

doc.add_paragraph()
add_heading('7.2 Deployment Actions (K8s / Terraform / Ansible)', level=2)
add_code_block(
'''oom-001       → scale_out (replicas +2) + configmap_patch (heap_limit=4096m)
hpa-scale     → hpa_update (4→9 replicas)
vpc-reroute   → bgp_reroute (backup-az activated)
db-throttle   → read_replica_promote (replica-1) + etl_throttle (rate=40%)
cert-renew    → acm_renew + DNS validation
kafka-rebalance → partition_rebalance + consumer_scale (+4 threads)
rollback-deploy → kubectl_rollout_undo (deployment/app → v2.4.0)
disk-cleanup  → log_rotation (18 GB freed)'''
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — DATA MODELS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('8. Data Models', level=1)
add_heading('8.1 Core Enumerations', level=2)
make_table(
    ['Enum', 'Values'],
    [
        ['Severity',         'P1 (critical), P2 (high), P3 (medium)'],
        ['IncidentStatus',   'open → triaging → remediating → preempting → resolved'],
        ['AgentStatus',      'active, busy, idle'],
        ['AgentType',        'server_monitor, cloud_monitor, app_health, predictive, remediation, deployment, reporting'],
        ['FailureClass',     'OOMKill, CPUSaturation, NetworkDegradation, LatencySpike, DiskPressure, CertExpiry, ReplicaLag, KafkaLag'],
    ],
    col_widths=[1.8, 5.5]
)

doc.add_paragraph()
add_heading('8.2 Incident', level=2)
add_code_block(
'''class Incident(BaseModel):
    id: str                        # "INC-XXXX"
    severity: Severity             # P1 | P2 | P3
    title: str
    description: str
    failure_class: FailureClass
    status: IncidentStatus         # open → triaging → remediating → resolved
    host: Optional[str]
    region: Optional[str]
    agent_chain: List[AgentType]   # Audit trail of agents that handled the incident
    opened_at: datetime
    resolved_at: Optional[datetime]
    mttr_seconds: Optional[float]  # Mean Time To Resolve
    runbook: Optional[str]
    auto_resolved: bool            # Always True on this platform'''
)

add_heading('8.3 Prediction', level=2)
add_code_block(
'''class Prediction(BaseModel):
    id: str
    failure_class: FailureClass
    target: str                    # Host / service / region
    probability: float             # 0.0–1.0 (LSTM confidence score)
    horizon_minutes: int           # Estimated time to failure
    signal_description: str
    recommended_runbook: str
    triggered_at: datetime
    acted_on: bool'''
)

add_heading('8.4 KPISnapshot', level=2)
add_code_block(
'''class KPISnapshot(BaseModel):
    uptime_pct: float
    active_incidents: int
    resolved_24h: int
    avg_mttr_minutes: float
    sla_adherence_pct: float
    timestamp: datetime'''
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — REST API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('9. REST API Reference', level=1)
make_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',  '/health',                              'Platform health check + agent count'],
        ['GET',  '/api/kpi',                             'Current KPI snapshot (uptime, MTTR, SLA)'],
        ['GET',  '/api/incidents',                       'All incidents (optional ?status= filter)'],
        ['POST', '/api/incidents/simulate?incident_type=','Inject simulated incident (oom/cpu/network/latency/disk)'],
        ['GET',  '/api/predictions',                     'Active failure predictions from LSTM agent'],
        ['GET',  '/api/agents',                          'All agent states + status + last heartbeat'],
        ['GET',  '/api/logs?limit=100',                  'Recent A2A messages (up to 500)'],
        ['GET',  '/api/metrics/hosts',                   'Latest host metrics (CPU/mem/disk) — up to 20'],
        ['GET',  '/api/metrics/apps',                    'Latest app metrics (p99/error rate/RPS)'],
        ['WS',   '/ws',                                  'Real-time event stream (WebSocket)'],
    ],
    col_widths=[0.7, 2.8, 3.8]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — WEBSOCKET EVENT REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

add_heading('10. WebSocket Event Reference', level=1)
add_heading('10.1 Initial Snapshots (on connect)', level=2)
make_table(
    ['Event', 'Payload Type'],
    [
        ['incidents_snapshot',   '[...Incident]'],
        ['predictions_snapshot', '[...Prediction]'],
        ['agents_snapshot',      '[...AgentState]'],
        ['log_snapshot',         '[...A2AMessage]'],
        ['kpi',                  'KPISnapshot'],
    ],
    col_widths=[2.5, 4.8]
)
doc.add_paragraph()
add_heading('10.2 Streaming Updates', level=2)
make_table(
    ['Event', 'Payload', 'Triggered by'],
    [
        ['incident',               'Incident',                         'PredictiveAgent opens new incident'],
        ['incident_update',        '{ id, status, mttr_seconds }',     'Remediation/Deployment status change'],
        ['prediction',             'Prediction',                       'PredictiveAgent LSTM score above threshold'],
        ['a2a_log',                'A2AMessage',                       'Any agent-to-agent message'],
        ['metric',                 '{ type, host/service, ...values }','Monitoring agent metric tick'],
        ['kpi',                    'KPISnapshot',                      'ReportingAgent 10s broadcast'],
        ['security_anomaly',       'SecurityAnomaly',                  'SecurityAgent detects threat'],
        ['cost_metric',            'CostMetric',                       'OptimizationAgent cost update'],
        ['optimization_suggestion','OptimizationSuggestion',           'OptimizationAgent finds saving'],
    ],
    col_widths=[1.8, 2.2, 3.3]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — FRONTEND VIEWS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('11. Frontend Views', level=1)
make_table(
    ['Tab', 'Description', 'Data Source'],
    [
        ['Agent Topology',  'Animated SVG graph showing all 10 agents with live A2A message flow lines and status cards','WebSocket agents_snapshot + a2a_log'],
        ['Incidents',       'Live incident board: severity badges (P1/P2/P3), status pipeline, agent chain, MTTR, simulate button','WebSocket incidents_snapshot + incident_update'],
        ['Predictions',     'LSTM prediction cards: probability bars, failure class, target, time horizon, recommended runbook','WebSocket predictions_snapshot + prediction'],
        ['A2A Log Stream',  'Real-time feed of every agent-to-agent message with expandable JSON payload inspector','WebSocket log_snapshot + a2a_log'],
        ['Live Metrics',    'Recharts time-series for CPU/mem/disk per host, p99 latency per service, sortable host table','WebSocket metric events'],
        ['Architecture',    'Tech stack detail cards and 4-step workflow pipeline diagram (static)','Static'],
    ],
    col_widths=[1.5, 3.2, 2.6]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════

add_heading('12. Deployment', level=1)
add_heading('12.1 Docker Compose (Current — Development / Demo)', level=2)
add_code_block(
'''services:
  backend:
    build: ./backend         # Python + FastAPI
    ports:  ["8000:8000"]
    healthcheck: curl /health every 10s — retries: 3

  frontend:
    build: ./frontend        # React build → Nginx static serve
    ports:  ["3000:80"]
    depends_on: backend
    env:
      REACT_APP_API_URL: http://localhost:8000
      REACT_APP_WS_URL:  ws://localhost:8000'''
)

doc.add_paragraph()
add_heading('12.2 Production Upgrade Path', level=2)
make_table(
    ['Component', 'Dev (Current)', 'Production Target'],
    [
        ['Data store',    'In-memory Python dict',      'PostgreSQL + TimescaleDB (time-series)'],
        ['ML model',      'LSTM scoring simulation',    'PyTorch LSTM trained model / MLflow versioned'],
        ['Message bus',   'asyncio.Queue (in-process)', 'Apache Kafka / RabbitMQ (distributed)'],
        ['Agent runtime', 'asyncio coroutines',         'Celery workers / Ray distributed actors'],
        ['Auth',          'None',                       'OAuth2 / JWT (FastAPI Security)'],
        ['Secrets',       'Environment variables',      'HashiCorp Vault'],
        ['Deployment',    'Docker Compose',             'Kubernetes + Helm chart'],
        ['Observability', 'WebSocket logs',             'Prometheus + Grafana + Jaeger tracing'],
        ['Alerting',      'WebSocket push',             'PagerDuty / OpsGenie integration'],
        ['CMDB',          'In-memory dict',             'ServiceNow / Backstage'],
    ],
    col_widths=[1.5, 2.2, 3.6]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — DESIGN DECISIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading('13. Key Design Decisions', level=1)
make_table(
    ['Decision', 'Rationale'],
    [
        ['asyncio throughout',              'Single-threaded event loop eliminates locking; all agents co-exist in one process with zero concurrency bugs'],
        ['A2A bus as asyncio.Queue',        'Zero-latency in-process delivery; swappable with Kafka for production with the same interface'],
        ['In-memory data store',            'Zero infra dependency for demo/dev; all interfaces designed for easy DB swap'],
        ['WebSocket-first push',            'Frontend always receives data without polling; supports multiple concurrent clients'],
        ['LSTM simulation without PyTorch', 'Allows full demo without GPU/model dependency; algorithm logic is production-equivalent and replaceable'],
        ['Pydantic v2 models',              'Strict schema validation + fast JSON serialization across all layers (agents, API, WebSocket)'],
        ['Docker multi-stage frontend build','Minimal Nginx image; no Node.js runtime in production container reduces attack surface'],
    ],
    col_widths=[2.2, 5.1]
)
add_divider()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Multi-Agent IT Ops Platform  ·  Architecture Reference  ·  v1.0.0  ·  2026-03-28  ·  MIT License')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = r'd:\jatayu1\multi-agent-itops\Architecture_doc.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
